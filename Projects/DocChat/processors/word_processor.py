"""
Word Processor - обработка Word документов (.docx)
"""

from docx import Document
from io import BytesIO
from typing import Dict


def process_word(file_bytes: bytes, filename: str) -> dict:
    """
    Обработка Word документа

    Args:
        file_bytes: байты Word файла
        filename: имя файла

    Returns:
        dict с результатами обработки
    """
    result = {
        "filename": filename,
        "type": "word",
        "text": None,
        "paragraphs": 0,
        "tables": 0
    }

    try:
        doc = Document(BytesIO(file_bytes))
        text_parts = []

        # Извлекаем параграфы
        for para in doc.paragraphs:
            if para.text.strip():
                # Определяем стиль (заголовок или обычный текст)
                style_name = para.style.name if para.style else ""

                if "Heading" in style_name or "Заголовок" in style_name:
                    level = 1
                    if "1" in style_name:
                        level = 1
                    elif "2" in style_name:
                        level = 2
                    elif "3" in style_name:
                        level = 3
                    text_parts.append(f"{'#' * level} {para.text}")
                else:
                    text_parts.append(para.text)

                result["paragraphs"] += 1

        # Извлекаем таблицы
        for table_idx, table in enumerate(doc.tables, 1):
            result["tables"] += 1
            text_parts.append(f"\n--- Таблица {table_idx} ---")

            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(" | ".join(row_data))

            text_parts.append("\n".join(table_data))

        result["text"] = "\n\n".join(text_parts)

    except Exception as e:
        result["text"] = f"[Ошибка обработки Word: {e}]"

    return result
