#!/usr/bin/env python3
"""
MCP сервер для обработки документов через Claude API
Поддержка: PDF, DOCX, XLSX, PPTX, изображения
"""

import asyncio
import base64
import json
import os
import sys
from pathlib import Path
from typing import Any

import anthropic
from mcp.server import Server
from mcp.server.stdio import stdio_server
from mcp.types import Tool, TextContent

# Импорты для обработки документов
try:
    import PyPDF2
    from pdf2image import convert_from_path
except ImportError:
    PyPDF2 = None
    convert_from_path = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    from openpyxl import load_workbook
except ImportError:
    load_workbook = None

try:
    from PIL import Image
    import io
except ImportError:
    Image = None
    io = None


# Получаем API ключ из переменной окружения
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY")

if not ANTHROPIC_API_KEY:
    print("ОШИБКА: Не найден ANTHROPIC_API_KEY в переменных окружения", file=sys.stderr)
    sys.exit(1)

# Инициализируем Claude клиент
claude_client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

# Создаем MCP сервер
app = Server("document-processor")


def encode_image_to_base64(image_path: str) -> tuple[str, str]:
    """Кодирует изображение в base64 и определяет MIME тип"""
    path = Path(image_path)

    # Определяем MIME тип
    extension = path.suffix.lower()
    mime_types = {
        '.jpg': 'image/jpeg',
        '.jpeg': 'image/jpeg',
        '.png': 'image/png',
        '.gif': 'image/gif',
        '.webp': 'image/webp'
    }
    mime_type = mime_types.get(extension, 'image/jpeg')

    # Читаем и кодируем файл
    with open(image_path, 'rb') as f:
        image_data = base64.standard_b64encode(f.read()).decode('utf-8')

    return image_data, mime_type


def pdf_to_images(pdf_path: str) -> list[str]:
    """Конвертирует PDF в изображения для отправки в Claude"""
    if not convert_from_path:
        raise ImportError("Установите pdf2image: pip install pdf2image")

    images = convert_from_path(pdf_path, dpi=200)
    image_paths = []

    temp_dir = Path(pdf_path).parent / "temp_pdf_images"
    temp_dir.mkdir(exist_ok=True)

    for i, image in enumerate(images):
        image_path = temp_dir / f"page_{i+1}.png"
        image.save(image_path, 'PNG')
        image_paths.append(str(image_path))

    return image_paths


def extract_text_from_docx(docx_path: str) -> str:
    """Извлекает текст из DOCX"""
    if not DocxDocument:
        raise ImportError("Установите python-docx: pip install python-docx")

    doc = DocxDocument(docx_path)
    text_parts = []

    for paragraph in doc.paragraphs:
        text_parts.append(paragraph.text)

    # Извлекаем текст из таблиц
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            text_parts.append(row_text)

    return "\n".join(text_parts)


def extract_text_from_xlsx(xlsx_path: str) -> str:
    """Извлекает данные из Excel в markdown таблицу"""
    if not load_workbook:
        raise ImportError("Установите openpyxl: pip install openpyxl")

    workbook = load_workbook(xlsx_path, data_only=True)
    result_parts = []

    for sheet_name in workbook.sheetnames:
        sheet = workbook[sheet_name]
        result_parts.append(f"## {sheet_name}\n")

        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            continue

        # Создаем markdown таблицу
        max_cols = max(len(row) for row in rows if row)

        for i, row in enumerate(rows[:100]):  # Ограничиваем 100 строками
            if not any(row):
                continue

            row_data = [str(cell) if cell is not None else "" for cell in row[:max_cols]]
            result_parts.append("| " + " | ".join(row_data) + " |")

            # Добавляем разделитель после заголовка
            if i == 0:
                result_parts.append("| " + " | ".join(["---"] * len(row_data)) + " |")

        result_parts.append("\n")

    return "\n".join(result_parts)


async def process_with_claude(content: str | list, prompt: str = None) -> str:
    """Обрабатывает контент через Claude API"""

    # Формируем сообщение
    if isinstance(content, str):
        # Текстовый контент
        messages = [{
            "role": "user",
            "content": prompt or "Извлеки и структурируй весь текст из этого документа, сохраняя форматирование и таблицы:"
        }, {
            "role": "user",
            "content": content
        }]
    else:
        # Изображения
        user_message = []
        if prompt:
            user_message.append({
                "type": "text",
                "text": prompt
            })
        user_message.extend(content)

        messages = [{
            "role": "user",
            "content": user_message
        }]

    # Отправляем запрос к Claude 4.5 Sonnet (последняя версия, сентябрь 2025)
    response = claude_client.messages.create(
        model="claude-sonnet-4-5",  # Claude 4.5 Sonnet - лучшая модель для агентов, кодинга, документов
        max_tokens=64000,  # Claude 4.5 поддерживает до 64K выходных токенов
        messages=messages
    )

    return response.content[0].text


@app.list_tools()
async def list_tools() -> list[Tool]:
    """Список доступных инструментов"""
    return [
        Tool(
            name="process_document",
            description=(
                "Обрабатывает документ (PDF, DOCX, XLSX, PPTX, изображение) через Claude API. "
                "Возвращает распознанный текст с сохранением структуры, таблиц и форматирования. "
                "Качество распознавания идентично веб-версии Claude."
            ),
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "Полный путь к файлу для обработки"
                    },
                    "prompt": {
                        "type": "string",
                        "description": "Дополнительная инструкция для обработки (опционально). Например: 'Извлеки только таблицы' или 'Сделай краткое резюме'"
                    }
                },
                "required": ["file_path"]
            }
        )
    ]


@app.call_tool()
async def call_tool(name: str, arguments: Any) -> list[TextContent]:
    """Обработка вызова инструмента"""

    if name != "process_document":
        raise ValueError(f"Неизвестный инструмент: {name}")

    file_path = arguments.get("file_path")
    prompt = arguments.get("prompt")

    if not file_path:
        raise ValueError("Не указан file_path")

    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Файл не найден: {file_path}")

    extension = path.suffix.lower()

    try:
        # Обработка изображений
        if extension in ['.jpg', '.jpeg', '.png', '.gif', '.webp']:
            image_data, mime_type = encode_image_to_base64(file_path)

            content = [{
                "type": "image",
                "source": {
                    "type": "base64",
                    "media_type": mime_type,
                    "data": image_data
                }
            }]

            result = await process_with_claude(content, prompt)

        # Обработка PDF
        elif extension == '.pdf':
            # Конвертируем PDF в изображения
            image_paths = pdf_to_images(file_path)

            content = []
            for img_path in image_paths:
                image_data, mime_type = encode_image_to_base64(img_path)
                content.append({
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": mime_type,
                        "data": image_data
                    }
                })

            result = await process_with_claude(content, prompt)

            # Очищаем временные файлы
            for img_path in image_paths:
                Path(img_path).unlink()
            temp_dir = Path(image_paths[0]).parent
            if temp_dir.exists() and not list(temp_dir.iterdir()):
                temp_dir.rmdir()

        # Обработка DOCX
        elif extension in ['.docx', '.doc']:
            text = extract_text_from_docx(file_path)
            result = await process_with_claude(text, prompt)

        # Обработка Excel
        elif extension in ['.xlsx', '.xls']:
            text = extract_text_from_xlsx(file_path)
            result = await process_with_claude(text, prompt)

        else:
            raise ValueError(f"Неподдерживаемый формат файла: {extension}")

        return [TextContent(
            type="text",
            text=f"=== Результат обработки файла: {path.name} ===\n\n{result}"
        )]

    except Exception as e:
        return [TextContent(
            type="text",
            text=f"Ошибка при обработке файла: {str(e)}"
        )]


async def main():
    """Запуск MCP сервера"""
    async with stdio_server() as (read_stream, write_stream):
        await app.run(
            read_stream,
            write_stream,
            app.create_initialization_options()
        )


if __name__ == "__main__":
    asyncio.run(main())
