#!/usr/bin/env python3
"""
Скрипт для извлечения изображений из .docx файла (Google Docs экспорт).

Использование:
    python extract_images_from_docx.py input.docx [output_folder]

Требования:
    pip install python-docx pillow
"""

import sys
import os
from pathlib import Path


def check_dependencies():
    """Проверка установленных зависимостей."""
    missing = []

    try:
        from docx import Document
    except ImportError:
        missing.append("python-docx")

    try:
        from PIL import Image
    except ImportError:
        missing.append("pillow")

    if missing:
        print(f"Отсутствуют библиотеки: {', '.join(missing)}")
        print(f"Установите их командой: pip install {' '.join(missing)}")
        return False
    return True


def extract_images(docx_path: str, output_folder: str = None) -> list:
    """
    Извлекает изображения из .docx файла.

    Args:
        docx_path: Путь к .docx файлу
        output_folder: Папка для сохранения (по умолчанию: images/ рядом с файлом)

    Returns:
        Список путей к извлечённым изображениям
    """
    from docx import Document
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from PIL import Image
    import io

    docx_path = Path(docx_path)

    if not docx_path.exists():
        print(f"Файл не найден: {docx_path}")
        return []

    # Определяем папку для вывода
    if output_folder is None:
        output_folder = docx_path.parent / "images"
    else:
        output_folder = Path(output_folder)

    output_folder.mkdir(parents=True, exist_ok=True)

    extracted_files = []
    image_counter = 1

    print(f"Открываю: {docx_path}")
    doc = Document(str(docx_path))

    # Извлекаем изображения из связей документа
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            try:
                image_data = rel.target_part.blob

                # Определяем формат изображения
                img_pil = Image.open(io.BytesIO(image_data))

                # Пропускаем слишком маленькие изображения
                if img_pil.width < 50 or img_pil.height < 50:
                    continue

                # Формируем имя файла
                filename = f"step{image_counter}_img.png"
                filepath = output_folder / filename

                # Конвертируем в PNG и сохраняем
                if img_pil.mode in ('RGBA', 'P'):
                    img_pil = img_pil.convert('RGB')
                img_pil.save(str(filepath), "PNG")

                extracted_files.append(str(filepath))
                print(f"  Сохранено: {filename} ({img_pil.width}x{img_pil.height})")
                image_counter += 1

            except Exception as e:
                print(f"  Ошибка извлечения изображения: {e}")

    print(f"\nИтого извлечено: {len(extracted_files)} изображений")
    print(f"Сохранены в: {output_folder}")

    return extracted_files


def extract_text(docx_path: str) -> str:
    """
    Извлекает текст из .docx файла.

    Args:
        docx_path: Путь к .docx файлу

    Returns:
        Текст документа
    """
    from docx import Document

    doc = Document(docx_path)

    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Определяем стиль заголовка
            if para.style.name.startswith('Heading'):
                level = para.style.name.replace('Heading ', '')
                try:
                    level = int(level)
                    paragraphs.append('#' * level + ' ' + text)
                except ValueError:
                    paragraphs.append(text)
            else:
                paragraphs.append(text)

    return '\n\n'.join(paragraphs)


def generate_markdown_links(image_paths: list) -> str:
    """Генерирует Markdown ссылки на изображения."""
    lines = []
    for path in image_paths:
        filename = Path(path).name
        lines.append(f"![Описание](./images/{filename})")
    return "\n\n".join(lines)


def main():
    if len(sys.argv) < 2:
        print(__doc__)
        print("\nПример:")
        print("  python extract_images_from_docx.py document.docx")
        print("  python extract_images_from_docx.py document.docx ./output_images")
        print("\nДополнительные опции:")
        print("  --text    Также извлечь текст документа")
        sys.exit(1)

    if not check_dependencies():
        sys.exit(1)

    docx_path = sys.argv[1]
    output_folder = None
    extract_text_flag = False

    for arg in sys.argv[2:]:
        if arg == "--text":
            extract_text_flag = True
        else:
            output_folder = arg

    # Извлекаем изображения
    extracted = extract_images(docx_path, output_folder)

    if extracted:
        print("\n--- Markdown ссылки для вставки ---\n")
        print(generate_markdown_links(extracted))

    # Извлекаем текст, если запрошено
    if extract_text_flag:
        print("\n--- Текст документа (Markdown) ---\n")
        text = extract_text(docx_path)
        print(text)

        # Сохраняем текст в файл
        docx_path = Path(docx_path)
        md_path = docx_path.with_suffix('.md')
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(text)
        print(f"\nТекст сохранён в: {md_path}")


if __name__ == "__main__":
    main()
